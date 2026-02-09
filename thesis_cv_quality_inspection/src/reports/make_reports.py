import json
import shutil
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import onnxruntime as ort
import pandas as pd
import seaborn as sns
from PIL import Image
from docx import Document
from docx.shared import Inches
from sklearn.metrics import confusion_matrix
from ultralytics import YOLO

from src.utils.config import load_config


def _load_jsons(pattern: str):
    files = Path("outputs/metrics").glob(pattern)
    data = []
    for f in files:
        try:
            payload = json.loads(f.read_text(encoding="utf-8"))
            if isinstance(payload, dict):
                payload["_source_file"] = f.name
                payload["_source_mtime"] = f.stat().st_mtime
            data.append(payload)
        except Exception:
            continue
    return pd.DataFrame(data)


def _load_history(pattern: str):
    files = Path("outputs/history").glob(pattern)
    rows = []
    for f in files:
        try:
            payload = json.loads(f.read_text(encoding="utf-8"))
            parts = f.stem.split("__")
            dataset = parts[1] if len(parts) > 1 else "unknown"
            model = parts[2] if len(parts) > 2 else f.stem
            for i, ep in enumerate(payload.get("epoch", [])):
                rows.append(
                    {
                        "dataset": dataset,
                        "model": model,
                        "epoch": ep,
                        "train_loss": payload.get("train_loss", [None] * len(payload.get("epoch", [])))[i],
                        "train_accuracy": payload.get("train_accuracy", [None] * len(payload.get("epoch", [])))[i],
                        "val_loss": payload.get("val_loss", [None] * len(payload.get("epoch", [])))[i],
                        "val_accuracy": payload.get("val_accuracy", [None] * len(payload.get("epoch", [])))[i],
                    }
                )
        except Exception:
            continue
    return pd.DataFrame(rows)


def _save_table(df, name):
    out_csv = Path("docs/tables") / name
    out_csv.parent.mkdir(parents=True, exist_ok=True)
    df.to_csv(out_csv, index=False)
    out_xlsx = Path("docs/defense_pack") / name.replace(".csv", ".xlsx")
    out_xlsx.parent.mkdir(parents=True, exist_ok=True)
    df.to_excel(out_xlsx, index=False)


def _dedupe_latest(df: pd.DataFrame, key_cols):
    if df.empty or "_source_mtime" not in df.columns:
        return df
    keys = [c for c in key_cols if c in df.columns]
    if not keys:
        return df
    return df.sort_values("_source_mtime").drop_duplicates(keys, keep="last").reset_index(drop=True)


def _drop_internal_cols(df: pd.DataFrame):
    if df.empty:
        return df
    cols = [c for c in df.columns if not c.startswith("_source_")]
    return df[cols].copy()


def _clean_figures_dir():
    figs_dir = Path("docs/figures")
    figs_dir.mkdir(parents=True, exist_ok=True)
    generated_patterns = [
        "FIG_Classification_Accuracy_Comparison.png",
        "FIG_Classification_Accuracy_Heatmap.png",
        "FIG_Classification_Training_Dynamics.png",
        "FIG_Classification_Loss_*.png",
        "FIG_Classification_ValAcc_*.png",
        "FIG_Detection_mAP_Heatmap.png",
        "FIG_Tradeoff_mAP_vs_Latency.png",
        "FIG_YOLO_mAP_Comparison.png",
        "FIG_Detection_Eval_Coverage.png",
        "FIG_YOLO_results_*.png",
        "FIG_YOLO_confusion_matrix_*.png",
        "FIG_YOLO_confusion_matrix_normalized_*.png",
        "FIG_YOLO_BoxPR_curve_*.png",
        "FIG_YOLO_BoxP_curve_*.png",
        "FIG_YOLO_BoxR_curve_*.png",
        "FIG_YOLO_BoxF1_curve_*.png",
        "FIG_YOLO_Loss_*.png",
        "FIG_YOLO_mAPCurve_*.png",
        "FIG_Classification_Best_ConfusionMatrix_*.png",
        "FIG_YOLO_Best_ConfusionMatrix_*.png",
    ]
    for pattern in generated_patterns:
        for p in figs_dir.glob(pattern):
            p.unlink()


def _plot_bar(series, ylabel, title, out_path):
    if series.empty:
        return
    ax = series.plot(kind="bar")
    ax.set_ylabel(ylabel)
    ax.set_title(title)
    plt.tight_layout()
    plt.savefig(out_path, dpi=300)
    plt.close()


def _plot_heatmap(df, index_col, col_col, value_col, title, out_path):
    piv = df.pivot_table(index=index_col, columns=col_col, values=value_col, aggfunc="mean")
    if piv.empty:
        return
    plt.figure(figsize=(9, 5))
    sns.heatmap(piv, annot=True, fmt=".3f", cmap="viridis")
    plt.title(title)
    plt.tight_layout()
    plt.savefig(out_path, dpi=300)
    plt.close()


def _plot_notice(title: str, message: str, out_path: Path):
    plt.figure(figsize=(9, 4))
    plt.axis("off")
    plt.title(title)
    plt.text(0.5, 0.5, message, ha="center", va="center", wrap=True)
    plt.tight_layout()
    plt.savefig(out_path, dpi=300)
    plt.close()


def _configured_detection_models(cfg):
    models = cfg.get("detector", {}).get("models", [])
    keep = set()
    for model in models:
        keep.add(model)
        keep.add(f"{model}_pruned")
    return keep


def _prepare_detection_for_reporting(det_df: pd.DataFrame, cfg):
    if det_df.empty:
        return det_df
    df = det_df.copy()
    if "runtime" not in df.columns:
        df["runtime"] = "onnx"
    if "onnx_variant" not in df.columns:
        df["onnx_variant"] = ""
    keep = _configured_detection_models(cfg)
    if keep:
        df = df[df["model"].isin(keep)].copy()
    if df.empty:
        return df
    for col in ["map50", "precision", "recall"]:
        if col not in df.columns:
            df[col] = np.nan
    no_det = (df["map50"].fillna(0) <= 0) & (df["precision"].fillna(0) <= 0) & (df["recall"].fillna(0) <= 0)
    df["eval_status"] = np.where(no_det, "no_detections", "ok")
    df["map50_plot"] = df["map50"].where(df["eval_status"] == "ok", np.nan)
    return df


def _plot_detection_coverage(det_df: pd.DataFrame, out_path: Path):
    if det_df.empty:
        return
    summary = det_df.groupby(["dataset", "model"])["eval_status"].apply(
        lambda s: "ok" if (s == "ok").any() else "no_detections"
    )
    grid = summary.unstack("model")
    if grid.empty:
        return
    numeric = grid.apply(lambda c: c.map({"ok": 1.0, "no_detections": 0.0}))
    annot = grid.replace({"ok": "OK", "no_detections": "NO DET"})
    plt.figure(figsize=(9, 5))
    sns.heatmap(numeric, annot=annot, fmt="", cmap="viridis", vmin=0, vmax=1, cbar=False)
    plt.title("Detection Evaluation Coverage (OK vs No Detections)")
    plt.tight_layout()
    plt.savefig(out_path, dpi=300)
    plt.close()


def _normalize_series(series: pd.Series, higher_is_better: bool):
    s = pd.to_numeric(series, errors="coerce")
    if s.notna().sum() == 0:
        return pd.Series(np.zeros(len(series)), index=series.index, dtype=float)
    mn = float(s.min())
    mx = float(s.max())
    if mx <= mn:
        return pd.Series(np.ones(len(series)), index=series.index, dtype=float)
    if higher_is_better:
        return (s - mn) / (mx - mn)
    return (mx - s) / (mx - mn)


def _ranking_weights(cfg):
    weights = cfg.get("analysis", {}).get("ranking_weights", {})
    w_acc = float(weights.get("accuracy", 0.5))
    w_lat = float(weights.get("latency", 0.3))
    w_energy = float(weights.get("energy", 0.2))
    total = w_acc + w_lat + w_energy
    if total <= 0:
        return 0.5, 0.3, 0.2
    return w_acc / total, w_lat / total, w_energy / total


def _primary_classification_rows(class_df: pd.DataFrame, cfg):
    if class_df.empty:
        return class_df
    profile = cfg.get("analysis", {}).get("ranking_profile", "EDGE_CPU_1T")
    df = class_df.copy()
    if "runtime" in df.columns:
        onnx_rows = df[df["runtime"] == "onnx"].copy()
        if not onnx_rows.empty:
            df = onnx_rows
    if "onnx_variant" in df.columns:
        fp_rows = df[df["onnx_variant"].fillna("") == "fp32"].copy()
        if not fp_rows.empty:
            df = fp_rows
    prof_rows = df[df["profile"] == profile].copy()
    if not prof_rows.empty:
        df = prof_rows
    return df


def _primary_detection_rows(det_df: pd.DataFrame, cfg):
    if det_df.empty:
        return det_df
    profile = cfg.get("analysis", {}).get("ranking_profile", "EDGE_CPU_1T")
    df = det_df.copy()
    if "runtime" in df.columns:
        onnx_rows = df[df["runtime"] == "onnx"].copy()
        if not onnx_rows.empty:
            df = onnx_rows
    prof_rows = df[df["profile"] == profile].copy()
    if not prof_rows.empty:
        df = prof_rows
    return df


def _build_weighted_rankings(class_df: pd.DataFrame, det_df: pd.DataFrame, cfg):
    w_acc, w_lat, w_energy = _ranking_weights(cfg)
    ranking_rows = []

    cls_base = _primary_classification_rows(class_df, cfg)
    if not cls_base.empty:
        cls = cls_base.copy()
        cls["accuracy_norm"] = _normalize_series(cls["accuracy"], higher_is_better=True)
        cls["latency_norm"] = _normalize_series(cls["latency_ms"], higher_is_better=False)
        cls["energy_norm"] = _normalize_series(cls["energy_proxy_wh_per_1000"], higher_is_better=False)
        cls["weighted_score"] = (
            w_acc * cls["accuracy_norm"] + w_lat * cls["latency_norm"] + w_energy * cls["energy_norm"]
        )
        cls["task"] = "classification"
        ranking_rows.append(
            cls[
                [
                    "task",
                    "dataset",
                    "model",
                    "profile",
                    "runtime",
                    "onnx_variant",
                    "accuracy",
                    "latency_ms",
                    "energy_proxy_wh_per_1000",
                    "accuracy_norm",
                    "latency_norm",
                    "energy_norm",
                    "weighted_score",
                ]
            ]
        )

    det_base = _primary_detection_rows(det_df, cfg)
    if not det_base.empty:
        det = det_base.copy()
        det["accuracy_norm"] = _normalize_series(det["map50_plot"], higher_is_better=True)
        det["latency_norm"] = _normalize_series(det["latency_ms"], higher_is_better=False)
        det["energy_norm"] = _normalize_series(det["energy_proxy_wh_per_1000"], higher_is_better=False)
        det["weighted_score"] = (
            w_acc * det["accuracy_norm"] + w_lat * det["latency_norm"] + w_energy * det["energy_norm"]
        )
        det["task"] = "detection"
        det["accuracy"] = det["map50_plot"]
        ranking_rows.append(
            det[
                [
                    "task",
                    "dataset",
                    "model",
                    "profile",
                    "runtime",
                    "onnx_variant",
                    "accuracy",
                    "latency_ms",
                    "energy_proxy_wh_per_1000",
                    "accuracy_norm",
                    "latency_norm",
                    "energy_norm",
                    "weighted_score",
                ]
            ]
        )

    if not ranking_rows:
        return pd.DataFrame(), pd.DataFrame(), pd.DataFrame()
    ranking_df = pd.concat(ranking_rows, ignore_index=True)

    detection_rec = ranking_df[ranking_df["task"] == "detection"].copy()
    if not detection_rec.empty:
        idx = detection_rec.groupby("dataset")["weighted_score"].idxmax()
        detection_rec = detection_rec.loc[idx].sort_values(["dataset"]).reset_index(drop=True)
    classification_rec = ranking_df[ranking_df["task"] == "classification"].copy()
    if not classification_rec.empty:
        idx = classification_rec.groupby("dataset")["weighted_score"].idxmax()
        classification_rec = classification_rec.loc[idx].sort_values(["dataset"]).reset_index(drop=True)
    return ranking_df, detection_rec, classification_rec


def _flatten_corr(df: pd.DataFrame, cols, table_name: str):
    rows = []
    if df.empty:
        return pd.DataFrame(rows)
    existing = [c for c in cols if c in df.columns]
    if len(existing) < 2:
        return pd.DataFrame(rows)
    corr = df[existing].apply(pd.to_numeric, errors="coerce").corr()
    for a in existing:
        for b in existing:
            rows.append({"table": table_name, "metric_a": a, "metric_b": b, "pearson_r": float(corr.loc[a, b])})
    return pd.DataFrame(rows)


def _build_correlation_tables(class_df: pd.DataFrame, det_df: pd.DataFrame, cfg):
    cls_base = _primary_classification_rows(class_df, cfg)
    det_base = _primary_detection_rows(det_df, cfg)
    cls_corr = _flatten_corr(
        cls_base,
        ["accuracy", "latency_ms", "fps", "energy_proxy_wh_per_1000", "accuracy_gap_vs_pt"],
        "classification",
    )
    det_corr = _flatten_corr(
        det_base,
        ["map50_plot", "latency_ms", "fps", "energy_proxy_wh_per_1000"],
        "detection",
    )
    return cls_corr, det_corr


def _build_classification_parity_table(class_df: pd.DataFrame, cfg):
    if class_df.empty or "accuracy_gap_vs_pt" not in class_df.columns:
        return pd.DataFrame()
    df = _primary_classification_rows(class_df, cfg).copy()
    if df.empty:
        return pd.DataFrame()
    cols = [
        "dataset",
        "model",
        "profile",
        "runtime",
        "onnx_variant",
        "accuracy",
        "reference_pt_accuracy",
        "accuracy_gap_vs_pt",
        "class_index_source",
    ]
    existing = [c for c in cols if c in df.columns]
    max_gap = float(cfg.get("analysis", {}).get("max_accuracy_gap_vs_pt", 0.05))
    out = df[existing].copy()
    gap = pd.to_numeric(out["accuracy_gap_vs_pt"], errors="coerce")
    out["parity_status"] = np.where(gap.abs() <= max_gap, "ok", "review")
    return out


def _build_runtime_coverage(class_df: pd.DataFrame, det_df: pd.DataFrame):
    rows = []
    for task_name, df in [("classification", class_df), ("detection", det_df)]:
        if df.empty:
            continue
        cols = [c for c in ["dataset", "model", "runtime", "onnx_variant", "profile"] if c in df.columns]
        if not cols:
            continue
        tmp = df[cols].copy()
        if "onnx_variant" not in tmp.columns:
            tmp["onnx_variant"] = ""
        grouped = (
            tmp.groupby(["dataset", "model", "runtime", "onnx_variant"], dropna=False)["profile"]
            .nunique()
            .reset_index(name="profiles_covered")
        )
        grouped.insert(0, "task", task_name)
        rows.append(grouped)
    if not rows:
        return pd.DataFrame()
    return pd.concat(rows, ignore_index=True)


def _resolve_cls_onnx_candidates(dataset: str, model_label: str):
    base = model_label
    suffix = ""
    if model_label.endswith("_pruned"):
        base = model_label[: -len("_pruned")]
        suffix = "_pruned"
    int8 = Path("outputs/models_onnx") / f"{dataset}__{base}{suffix}_int8.onnx"
    fp32 = Path("outputs/models_onnx") / f"{dataset}__{base}{suffix}.onnx"
    candidates = []
    if fp32.exists():
        candidates.append(fp32)
    if int8.exists():
        candidates.append(int8)
    return candidates


def _load_cls_test_samples(processed_dir: Path):
    labels_path = processed_dir / "labels_classification.csv"
    splits_path = processed_dir / "splits.csv"
    if not labels_path.exists() or not splits_path.exists():
        return []
    labels = pd.read_csv(labels_path)
    splits = pd.read_csv(splits_path)
    test = labels.merge(splits, on="image_path")
    test = test[test["split"] == "test"]
    return test


def _infer_cls_confusion_matrix(dataset: str, model_label: str, cfg):
    dcfg = cfg["data"]["datasets"][dataset]
    class_names = list(dcfg["class_names"])
    processed_dir = Path(dcfg["processed_dir"])
    train_root = processed_dir / "images" / "train"
    if train_root.exists():
        inferred = sorted([p.name for p in train_root.iterdir() if p.is_dir()])
        if inferred:
            class_names = inferred
    class_to_idx = {c: i for i, c in enumerate(class_names)}
    candidates = _resolve_cls_onnx_candidates(dataset, model_label)
    if not candidates or not processed_dir.exists():
        return None, None
    test = _load_cls_test_samples(processed_dir)
    if test.empty:
        return None, None
    sess = None
    for onnx_path in candidates:
        try:
            sess = ort.InferenceSession(str(onnx_path), providers=["CPUExecutionProvider"])
            break
        except Exception:
            continue
    if sess is None:
        return None, None
    input_name = sess.get_inputs()[0].name
    input_size = cfg["classification"]["input_size"]

    y_true = []
    y_pred = []
    for _, row in test.iterrows():
        img_path = Path(row["image_path"])
        if not img_path.exists():
            continue
        img = Image.open(img_path).convert("RGB").resize((input_size, input_size))
        arr = np.array(img).astype(np.float32) / 255.0
        arr = np.transpose(arr, (2, 0, 1))[None, ...]
        out = sess.run(None, {input_name: arr})[0][0]
        pred = int(np.argmax(out))
        label = row["class"]
        if label not in class_to_idx:
            continue
        y_pred.append(pred)
        y_true.append(class_to_idx[label])

    if not y_true:
        return None, None
    cm = confusion_matrix(y_true, y_pred, labels=list(range(len(class_names))), normalize="true")
    return cm, class_names


def _pick_best_classification_models(class_df: pd.DataFrame):
    if class_df.empty:
        return pd.DataFrame()
    filt = class_df[(class_df["profile"] == "EDGE_CPU_1T") & (class_df["runtime"] == "onnx")].copy()
    if filt.empty:
        filt = class_df[class_df["runtime"] == "onnx"].copy()
    if filt.empty:
        return pd.DataFrame()
    idx = filt.groupby("dataset")["accuracy"].idxmax()
    return filt.loc[idx, ["dataset", "model", "accuracy"]].reset_index(drop=True)


def _pick_best_detection_models(det_df: pd.DataFrame):
    if det_df.empty:
        return pd.DataFrame()
    filt = det_df.copy()
    if "runtime" in filt.columns:
        onnx_rows = filt[filt["runtime"] == "onnx"].copy()
        if not onnx_rows.empty:
            filt = onnx_rows
    filt = filt[filt["profile"] == "EDGE_CPU_1T"].copy()
    if filt.empty:
        filt = det_df.copy()
    valid = filt[(filt["eval_status"] == "ok") & (filt["map50_plot"].notna())]
    if valid.empty:
        return pd.DataFrame()
    idx = valid.groupby("dataset")["map50_plot"].idxmax()
    return valid.loc[idx, ["dataset", "model", "map50"]].reset_index(drop=True)


def _copy_best_yolo_confusion_mats(best_det_df: pd.DataFrame, cfg):
    figs_dir = Path("docs/figures")
    copied_datasets = set()
    for _, row in best_det_df.iterrows():
        dataset = str(row["dataset"])
        copied_datasets.add(dataset)
        model_label = str(row["model"])
        base_model = model_label.replace("_pruned", "")
        run_dir = Path("outputs/logs") / f"{dataset}__{base_model}"
        if not run_dir.exists():
            alt = Path("runs/detect/outputs/logs") / f"{dataset}__{base_model}"
            if alt.exists():
                run_dir = alt
        norm = run_dir / "confusion_matrix_normalized.png"
        raw = run_dir / "confusion_matrix.png"
        src = norm if norm.exists() else raw
        if src.exists():
            shutil.copy2(src, figs_dir / f"FIG_YOLO_Best_ConfusionMatrix_{dataset}.png")

    enabled_det_datasets = [
        name for name, dcfg in cfg.get("data", {}).get("datasets", {}).items() if dcfg.get("enabled_detection", False)
    ]
    for dataset in enabled_det_datasets:
        if dataset in copied_datasets:
            continue
        _plot_notice(
            "Best Detection Confusion Matrix",
            f"No valid confusion matrix for {dataset}: all configured detection models produced no detections.",
            figs_dir / f"FIG_YOLO_Best_ConfusionMatrix_{dataset}.png",
        )


def make_essential_figures(class_df: pd.DataFrame, det_df: pd.DataFrame, hist_df: pd.DataFrame):
    figs_dir = Path("docs/figures")
    figs_dir.mkdir(parents=True, exist_ok=True)

    if not class_df.empty:
        _plot_bar(
            class_df.groupby("model")["accuracy"].mean().sort_values(ascending=False),
            "Accuracy",
            "Classification Accuracy (Mean Across Datasets/Profiles)",
            figs_dir / "FIG_Classification_Accuracy_Comparison.png",
        )
        cls_heat = class_df[(class_df["profile"] == "EDGE_CPU_1T") & (class_df["runtime"] == "onnx")]
        if cls_heat.empty:
            cls_heat = class_df[class_df["runtime"] == "onnx"]
        _plot_heatmap(
            cls_heat,
            "dataset",
            "model",
            "accuracy",
            "Classification Accuracy Heatmap",
            figs_dir / "FIG_Classification_Accuracy_Heatmap.png",
        )

    if not det_df.empty:
        det_plot_df = det_df.copy()
        if "runtime" in det_plot_df.columns:
            onnx_rows = det_plot_df[det_plot_df["runtime"] == "onnx"].copy()
            if not onnx_rows.empty:
                det_plot_df = onnx_rows
        det_bar = det_plot_df.groupby("model")["map50_plot"].mean().sort_values(ascending=False).dropna()
        if det_bar.empty:
            _plot_notice(
                "Detection mAP@0.5 (Mean Across Datasets/Profiles)",
                "No valid detection runs to plot. All configured models returned no detections.",
                figs_dir / "FIG_YOLO_mAP_Comparison.png",
            )
        else:
            _plot_bar(
                det_bar,
                "mAP@0.5",
                "Detection mAP@0.5 (Mean Across Datasets/Profiles)",
                figs_dir / "FIG_YOLO_mAP_Comparison.png",
            )
        det_heat = det_plot_df[det_plot_df["profile"] == "EDGE_CPU_1T"]
        if det_heat.empty:
            det_heat = det_plot_df
        det_heat = det_heat[det_heat["map50_plot"].notna()]
        if det_heat.empty:
            _plot_notice(
                "Detection mAP@0.5 Heatmap",
                "No valid detection heatmap: all configured models returned no detections.",
                figs_dir / "FIG_Detection_mAP_Heatmap.png",
            )
        else:
            _plot_heatmap(
                det_heat,
                "dataset",
                "model",
                "map50_plot",
                "Detection mAP@0.5 Heatmap",
                figs_dir / "FIG_Detection_mAP_Heatmap.png",
            )

        det_scatter = det_plot_df[det_plot_df["map50_plot"].notna() & det_plot_df["latency_ms"].notna()]
        if det_scatter.empty:
            _plot_notice(
                "Detection Tradeoff: mAP@0.5 vs Latency",
                "No valid points to plot because all configured detection runs returned no detections.",
                figs_dir / "FIG_Tradeoff_mAP_vs_Latency.png",
            )
        else:
            plt.scatter(det_scatter["map50_plot"], det_scatter["latency_ms"])
            plt.xlabel("mAP@0.5")
            plt.ylabel("Latency (ms)")
            plt.title("Detection Tradeoff: mAP@0.5 vs Latency")
            plt.tight_layout()
            plt.savefig(figs_dir / "FIG_Tradeoff_mAP_vs_Latency.png", dpi=300)
            plt.close()

        _plot_detection_coverage(det_plot_df, figs_dir / "FIG_Detection_Eval_Coverage.png")

    if not hist_df.empty:
        # One compact global trend figure for training dynamics.
        trend = hist_df.groupby("epoch")[["train_loss", "val_loss", "train_accuracy", "val_accuracy"]].mean()
        plt.figure(figsize=(9, 5))
        plt.plot(trend.index, trend["train_loss"], label="train_loss")
        plt.plot(trend.index, trend["val_loss"], label="val_loss")
        plt.plot(trend.index, trend["train_accuracy"], label="train_accuracy")
        plt.plot(trend.index, trend["val_accuracy"], label="val_accuracy")
        plt.xlabel("Epoch")
        plt.title("Classification Training Dynamics (Mean Across Models)")
        plt.legend()
        plt.tight_layout()
        plt.savefig(figs_dir / "FIG_Classification_Training_Dynamics.png", dpi=300)
        plt.close()


def _find_yolo_run_dir(dataset: str, model: str):
    candidates = [
        Path("outputs/logs") / f"{dataset}__{model}",
        Path("runs/detect/outputs/logs") / f"{dataset}__{model}",
    ]
    for c in candidates:
        if c.exists():
            return c
    return None


def _make_classification_per_model_figures(hist_df: pd.DataFrame):
    if hist_df.empty:
        return
    figs_dir = Path("docs/figures")
    grouped = hist_df.groupby(["dataset", "model"], dropna=False)
    for (dataset, model), g in grouped:
        if g.empty:
            continue
        g = g.sort_values("epoch")
        loss_cols = [c for c in ["train_loss", "val_loss"] if c in g.columns and g[c].notna().any()]
        if loss_cols:
            plt.figure(figsize=(8, 4))
            for col in loss_cols:
                plt.plot(g["epoch"], g[col], label=col)
            plt.xlabel("Epoch")
            plt.ylabel("Loss")
            plt.title(f"Classification Loss ({dataset}: {model})")
            plt.legend()
            plt.tight_layout()
            plt.savefig(figs_dir / f"FIG_Classification_Loss_{dataset}__{model}.png", dpi=300)
            plt.close()

        if "val_accuracy" in g.columns and g["val_accuracy"].notna().any():
            plt.figure(figsize=(8, 4))
            plt.plot(g["epoch"], g["val_accuracy"], label="val_accuracy")
            if "train_accuracy" in g.columns and g["train_accuracy"].notna().any():
                plt.plot(g["epoch"], g["train_accuracy"], label="train_accuracy")
            plt.xlabel("Epoch")
            plt.ylabel("Accuracy")
            plt.title(f"Classification Accuracy ({dataset}: {model})")
            plt.legend()
            plt.tight_layout()
            plt.savefig(figs_dir / f"FIG_Classification_ValAcc_{dataset}__{model}.png", dpi=300)
            plt.close()


def _copy_yolo_artifact_figures(cfg):
    figs_dir = Path("docs/figures")
    artifact_map = {
        "results.png": "FIG_YOLO_results_{dataset}__{model}.png",
        "confusion_matrix.png": "FIG_YOLO_confusion_matrix_{dataset}__{model}.png",
        "confusion_matrix_normalized.png": "FIG_YOLO_confusion_matrix_normalized_{dataset}__{model}.png",
        "BoxPR_curve.png": "FIG_YOLO_BoxPR_curve_{dataset}__{model}.png",
        "BoxP_curve.png": "FIG_YOLO_BoxP_curve_{dataset}__{model}.png",
        "BoxR_curve.png": "FIG_YOLO_BoxR_curve_{dataset}__{model}.png",
        "BoxF1_curve.png": "FIG_YOLO_BoxF1_curve_{dataset}__{model}.png",
    }
    for dataset, dcfg in cfg.get("data", {}).get("datasets", {}).items():
        if not dcfg.get("enabled_detection", False):
            continue
        for model in cfg.get("detector", {}).get("models", []):
            run_dir = _find_yolo_run_dir(dataset, model)
            if run_dir is None:
                continue
            for src_name, dst_tmpl in artifact_map.items():
                src = run_dir / src_name
                if src.exists():
                    shutil.copy2(src, figs_dir / dst_tmpl.format(dataset=dataset, model=model))
            results_csv = run_dir / "results.csv"
            if results_csv.exists():
                try:
                    df = pd.read_csv(results_csv)
                except Exception:
                    continue
                if "epoch" in df.columns:
                    if "train/box_loss" in df.columns or "val/box_loss" in df.columns:
                        plt.figure(figsize=(8, 4))
                        if "train/box_loss" in df.columns:
                            plt.plot(df["epoch"], df["train/box_loss"], label="train_box_loss")
                        if "val/box_loss" in df.columns:
                            plt.plot(df["epoch"], df["val/box_loss"], label="val_box_loss")
                        plt.xlabel("Epoch")
                        plt.ylabel("Loss")
                        plt.title(f"YOLO Box Loss ({dataset}: {model})")
                        plt.legend()
                        plt.tight_layout()
                        plt.savefig(figs_dir / f"FIG_YOLO_Loss_{dataset}__{model}.png", dpi=300)
                        plt.close()
                    if "metrics/mAP50(B)" in df.columns:
                        plt.figure(figsize=(8, 4))
                        plt.plot(df["epoch"], df["metrics/mAP50(B)"], label="mAP50")
                        if "metrics/mAP50-95(B)" in df.columns:
                            plt.plot(df["epoch"], df["metrics/mAP50-95(B)"], label="mAP50-95")
                        plt.xlabel("Epoch")
                        plt.ylabel("mAP")
                        plt.title(f"YOLO mAP ({dataset}: {model})")
                        plt.legend()
                        plt.tight_layout()
                        plt.savefig(figs_dir / f"FIG_YOLO_mAPCurve_{dataset}__{model}.png", dpi=300)
                        plt.close()


def make_best_confusion_mats(class_df: pd.DataFrame, det_df: pd.DataFrame, cfg):
    figs_dir = Path("docs/figures")
    best_cls = _pick_best_classification_models(class_df)
    best_det = _pick_best_detection_models(det_df)

    for _, row in best_cls.iterrows():
        dataset = str(row["dataset"])
        model_label = str(row["model"])
        cm, class_names = _infer_cls_confusion_matrix(dataset, model_label, cfg)
        if cm is None:
            continue
        plt.figure(figsize=(8, 6))
        sns.heatmap(cm, annot=True, fmt=".2f", cmap="magma", xticklabels=class_names, yticklabels=class_names)
        plt.xlabel("Predicted")
        plt.ylabel("True")
        plt.title(f"Best Classification Confusion Matrix ({dataset}: {model_label})")
        plt.tight_layout()
        plt.savefig(figs_dir / f"FIG_Classification_Best_ConfusionMatrix_{dataset}.png", dpi=300)
        plt.close()

    _copy_best_yolo_confusion_mats(best_det, cfg)


def make_appendix_docx():
    doc = Document()
    doc.add_heading("Results Appendix", level=1)
    doc.add_paragraph("Auto-generated appendix with essential tables and figures.")

    for fig in sorted(Path("docs/figures").glob("*.png")):
        doc.add_heading(fig.stem, level=2)
        doc.add_picture(str(fig), width=Inches(5.5))

    out_path = Path("docs/defense_pack/Results_Appendix.docx")
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def make_defense_pack():
    pack = Path("docs/defense_pack")
    pack.mkdir(parents=True, exist_ok=True)
    for p in Path("docs/tables").glob("*.csv"):
        shutil.copy2(p, pack / p.name)
    for p in Path("docs/figures").glob("*.png"):
        shutil.copy2(p, pack / p.name)
    (pack / "DEFENSE_SUMMARY.md").write_text(
        "Defense pack generated with essential charts, heatmaps, and best-model confusion matrices.\n",
        encoding="utf-8",
    )


def main():
    cfg = load_config()
    _clean_figures_dir()

    class_df = _load_jsons("classification_bench__*.json")
    class_df = _dedupe_latest(class_df, ["dataset", "model", "profile", "runtime", "onnx_variant"])
    det_df_raw = _load_jsons("detection_bench__*.json")
    det_df_raw = _dedupe_latest(det_df_raw, ["dataset", "model", "profile", "runtime", "onnx_variant"])
    det_df = _prepare_detection_for_reporting(det_df_raw, cfg)
    sim_df = _load_jsons("realtime_sim__*.json")
    sim_df = _dedupe_latest(sim_df, ["dataset", "model", "profile", "onnx_variant"])
    hist_df = _load_history("classification__*.json")

    if not class_df.empty:
        _save_table(_drop_internal_cols(class_df), "TABLE_Classification_Benchmark_All.csv")
    if not det_df.empty:
        _save_table(_drop_internal_cols(det_df), "TABLE_Detection_Benchmark_All.csv")
        det_trade = _primary_detection_rows(det_df, cfg)
        if det_trade.empty:
            det_trade = det_df.copy()
        trade_cols = ["dataset", "model", "map50", "latency_ms", "energy_proxy_wh_per_1000", "eval_status"]
        if "runtime" in det_trade.columns:
            trade_cols.insert(2, "runtime")
        if "onnx_variant" in det_trade.columns:
            trade_cols.insert(3, "onnx_variant")
        trade_cols = [c for c in trade_cols if c in det_trade.columns]
        _save_table(
            _drop_internal_cols(det_trade[trade_cols]),
            "TABLE_YOLO_Tradeoff_Accuracy_Latency_Energy.csv",
        )
    if not sim_df.empty:
        _save_table(_drop_internal_cols(sim_df), "TABLE_RealTime_Simulation.csv")
    if not hist_df.empty:
        _save_table(hist_df, "TABLE_Classification_TrainVal_History.csv")

    ranking_df, det_rec_df, cls_rec_df = _build_weighted_rankings(class_df, det_df, cfg)
    if not ranking_df.empty:
        _save_table(ranking_df, "TABLE_Model_Weighted_Ranking.csv")
    if not det_rec_df.empty:
        _save_table(det_rec_df, "TABLE_Detection_Recommended_Deployment.csv")
    if not cls_rec_df.empty:
        _save_table(cls_rec_df, "TABLE_Classification_Recommended_Deployment.csv")

    cls_corr_df, det_corr_df = _build_correlation_tables(class_df, det_df, cfg)
    corr_tables = []
    if not cls_corr_df.empty:
        corr_tables.append(cls_corr_df)
    if not det_corr_df.empty:
        corr_tables.append(det_corr_df)
    if corr_tables:
        _save_table(pd.concat(corr_tables, ignore_index=True), "TABLE_Correlation_Analysis.csv")

    parity_df = _build_classification_parity_table(class_df, cfg)
    if not parity_df.empty:
        _save_table(parity_df, "TABLE_Classification_Parity_Check.csv")

    runtime_cov_df = _build_runtime_coverage(class_df, det_df)
    if not runtime_cov_df.empty:
        _save_table(runtime_cov_df, "TABLE_Runtime_Coverage.csv")

    if not class_df.empty or not det_df.empty:
        model_sizes = []
        for df in [class_df, det_df]:
            if df.empty:
                continue
            model_sizes.append(df[["dataset", "model", "model_size_mb"]])
        if model_sizes:
            _save_table(pd.concat(model_sizes, ignore_index=True), "TABLE_Model_Size_Comparison.csv")

    make_essential_figures(class_df, det_df, hist_df)
    _make_classification_per_model_figures(hist_df)
    _copy_yolo_artifact_figures(cfg)
    make_best_confusion_mats(class_df, det_df, cfg)
    make_defense_pack()
    make_appendix_docx()


if __name__ == "__main__":
    main()
